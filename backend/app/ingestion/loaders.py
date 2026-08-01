from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader


PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_MIME_TYPES = {"text/plain", "text/markdown"}


class TextExtractionError(Exception):
    pass


@dataclass(frozen=True)
class ExtractedSection:
    text: str
    metadata: dict[str, Any]


@dataclass(frozen=True)
class ExtractedDocument:
    sections: list[ExtractedSection]

    @property
    def text(self) -> str:
        return "\n\n".join(section.text for section in self.sections)


def extract_text_from_file(file_path: str, mime_type: str) -> ExtractedDocument:
    path = Path(file_path)
    if not path.is_file():
        raise TextExtractionError(f"Document file does not exist: {path}")

    try:
        if mime_type in TEXT_MIME_TYPES:
            return _extract_text_file(path)
        if mime_type == PDF_MIME_TYPE:
            return _extract_pdf(path)
        if mime_type == DOCX_MIME_TYPE:
            return _extract_docx(path)
    except TextExtractionError:
        raise
    except Exception as exc:
        raise TextExtractionError(f"Could not extract text from {path.name}: {exc}") from exc

    raise TextExtractionError(f"Unsupported document MIME type: {mime_type}")


def _extract_text_file(path: Path) -> ExtractedDocument:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise TextExtractionError(
            f"{path.name} is not valid UTF-8 text"
        ) from exc
    if not text.strip():
        raise TextExtractionError(f"{path.name} contains no extractable text")
    return ExtractedDocument(
        sections=[ExtractedSection(text=text, metadata={"source": str(path)})]
    )


def _extract_pdf(path: Path) -> ExtractedDocument:
    reader = PdfReader(str(path))
    sections: list[ExtractedSection] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            sections.append(
                ExtractedSection(
                    text=text,
                    metadata={"source": str(path), "page_number": page_number},
                )
            )
    if not sections:
        raise TextExtractionError(
            f"{path.name} has no extractable text; scanned PDFs require OCR, "
            "which is not supported yet"
        )
    return ExtractedDocument(sections=sections)


def _extract_docx(path: Path) -> ExtractedDocument:
    document = DocxDocument(str(path))
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    if not paragraphs:
        raise TextExtractionError(f"{path.name} contains no extractable text")
    return ExtractedDocument(
        sections=[
            ExtractedSection(
                text="\n\n".join(paragraphs),
                metadata={"source": str(path)},
            )
        ]
    )
